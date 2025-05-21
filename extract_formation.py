from docx import Document
from io import BytesIO

def extract_section_by_header(file_bytes, target_header, all_known_headers):
    """
    Extracts paragraphs under a given header (e.g. 'Certifications') until the next known header.

    Args:
        file_bytes (bytes): The DOCX file in memory.
        target_header (str): Header to search for (e.g., 'Certifications').
        all_known_headers (list[str]): List of all other headers that may follow.

    Returns:
        list[str]: Paragraph texts under the target section.
    """
    target_header_lower = target_header.lower()
    known_headers_lower = [h.lower() for h in all_known_headers]

    with BytesIO(file_bytes) as stream:
        doc = Document(stream)
        capturing = False
        section_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            text_lower = text.lower()

            if not text:
                continue  # Skip empty lines

            if text_lower == target_header_lower:
                capturing = True
                continue

            if capturing and text_lower in known_headers_lower:
                break  # Stop when the next known header is found

            if capturing:
                section_content.append(text)

    return section_content


##################################

from docx import Document
from io import BytesIO

def extract_section_by_header(file_bytes, target_header, all_known_headers):
    """
    Extracts paragraphs under a given header (e.g. 'Certifications') until the next known header.

    Args:
        file_bytes (bytes): The DOCX file in memory.
        target_header (str): Header to search for (e.g., 'Certifications').
        all_known_headers (list[str]): List of all other headers that may follow.

    Returns:
        list[str]: Paragraph texts under the target section.
    """
    target_header_lower = target_header.lower()
    known_headers_lower = [h.lower() for h in all_known_headers]

    with BytesIO(file_bytes) as stream:
        doc = Document(stream)
        capturing = False
        section_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            text_lower = text.lower()

            if not text:
                continue  # Skip empty lines

            if text_lower == target_header_lower:
                capturing = True
                continue

            if capturing and text_lower in known_headers_lower:
                break  # Stop when the next known header is found

            if capturing:
                section_content.append(text)

    return section_content
🧪 Example Usage
python
Copier
Modifier
with open("CV-Gabarit-LGS-2023.docx", "rb") as f:
    file_bytes = f.read()

# Define all section headers that might follow "Certifications"
all_headers = [
    "Principaux domaines",
    "Formation académique",
    "Certifications",
    "Résumé des interventions",
    "Perfectionnement",
    "Langues parlées, écrites"
]

cert_section = extract_section_by_header(file_bytes, "Certifications", all_headers)

# Display extracted certification section
print("📝 Certifications Section:")
for line in cert_section:
    print(f"- {line}")
